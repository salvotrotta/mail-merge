"""
MAIL MERGE GUI
==============
Requisiti: Python 3.8+ (Tkinter incluso), LibreOffice per PDF
Librerie installate automaticamente: python-docx, openpyxl

Avvio: python mail_merge_gui.py
"""

import subprocess, sys, os, re, shutil, threading

# ── Log errori a file (utile quando compilato con PyInstaller) ───────────────
import traceback

def _log_crash(exc_type, exc_value, exc_tb):
    log_path = os.path.join(os.path.expanduser("~"), "MailMerge_error.log")
    with open(log_path, "a", encoding="utf-8") as f:
        import datetime
        f.write(f"\n{'='*60}\n{datetime.datetime.now()}\n")
        traceback.print_exception(exc_type, exc_value, exc_tb, file=f)
    try:
        import tkinter.messagebox as mb
        mb.showerror("Errore avvio",
                     f"Si è verificato un errore.\nDettagli in:\n{log_path}")
    except Exception:
        pass

sys.excepthook = _log_crash

for pkg in ["python-docx", "openpyxl"]:
    try:
        __import__(pkg.replace("-", "_").split(".")[0])
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"])

import openpyxl
from docx import Document
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

BG      = "#F8F7FF"
CARD    = "#FFFFFF"
ACCENT  = "#4F46E5"
ACCENT2 = "#6366F1"
TEXT    = "#1F2937"
MUTED   = "#6B7280"
BORDER  = "#E5E7EB"
SUCCESS = "#059669"
ERROR   = "#DC2626"
FONT    = "Segoe UI"

# ── Logica merge ──────────────────────────────────────────────────────────────
def trova_libreoffice():
    for p in [r"C:\Program Files\LibreOffice\program\soffice.exe",
              r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"]:
        if os.path.exists(p): return p
    return shutil.which("soffice")

def sostituisci_paragrafo(para, riga):
    testo = "".join(r.text for r in para.runs)
    if "{{" not in testo: return
    testo = re.sub(r'\{\{([^}]+)\}\}', lambda m: str(riga.get(m.group(1).strip(), m.group(0))), testo)
    if para.runs:
        para.runs[0].text = testo
        for r in para.runs[1:]: r.text = ""

def sostituisci_xml_raw(elemento, riga, colonne_valuta=None):
    """
    Sostituisce i segnaposto direttamente nell'XML dell'elemento,
    utile per raggiungere nodi non esposti da python-docx
    (indici, TOC, content controls, campi strutturati, ecc.)
    """
    colonne_valuta = colonne_valuta or set()
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Tutti i nodi <w:t> dentro l'elemento
    for t in elemento.iter("{%s}t" % ns):
        if t.text and "{{" in t.text:
            def repl(m):
                chiave = m.group(1).strip()
                val = riga.get(chiave, m.group(0))
                if chiave in colonne_valuta:
                    try:
                        n = float(str(val).replace(",", "."))
                        val = "€ {:.2f}".format(n).replace(",", "X").replace(".", ",").replace("X", ".")
                    except (ValueError, TypeError):
                        pass
                return str(val)
            t.text = re.sub(r'\{\{([^}]+)\}\}', repl, t.text)

def sostituisci_paragrafo(para, riga, colonne_valuta=None):
    colonne_valuta = colonne_valuta or set()
    testo = "".join(r.text for r in para.runs)
    if "{{" not in testo:
        return
    def repl(m):
        chiave = m.group(1).strip()
        val = riga.get(chiave, m.group(0))
        if chiave in colonne_valuta:
            try:
                n = float(str(val).replace(",", "."))
                val = "€ {:.2f}".format(n).replace(",", "X").replace(".", ",").replace("X", ".")
            except (ValueError, TypeError):
                pass
        return str(val)
    testo = re.sub(r'\{\{([^}]+)\}\}', repl, testo)
    if para.runs:
        para.runs[0].text = testo
        for r in para.runs[1:]: r.text = ""

def processa_documento(template_path, riga, colonne_valuta=None):
    colonne_valuta = colonne_valuta or set()
    doc = Document(template_path)

    # 1. Paragrafi normali
    for para in doc.paragraphs:
        sostituisci_paragrafo(para, riga, colonne_valuta)

    # 2. Tabelle
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    sostituisci_paragrafo(para, riga, colonne_valuta)

    # 3. Header e footer di ogni sezione
    for sez in doc.sections:
        for para in sez.header.paragraphs:
            sostituisci_paragrafo(para, riga, colonne_valuta)
        for para in sez.footer.paragraphs:
            sostituisci_paragrafo(para, riga, colonne_valuta)
        if sez.different_first_page_header_footer:
            for para in sez.first_page_header.paragraphs:
                sostituisci_paragrafo(para, riga, colonne_valuta)
            for para in sez.first_page_footer.paragraphs:
                sostituisci_paragrafo(para, riga, colonne_valuta)

    # 4. Indice / TOC / Content Controls / campi strutturati
    #    Agisce direttamente sull'XML del body per raggiungere
    #    nodi non esposti da python-docx
    sostituisci_xml_raw(doc.element.body, riga, colonne_valuta)

    # 5. XML di header e footer (per sicurezza)
    for sez in doc.sections:
        sostituisci_xml_raw(sez.header._element, riga, colonne_valuta)
        sostituisci_xml_raw(sez.footer._element, riga, colonne_valuta)
        if sez.different_first_page_header_footer:
            sostituisci_xml_raw(sez.first_page_header._element, riga, colonne_valuta)
            sostituisci_xml_raw(sez.first_page_footer._element, riga, colonne_valuta)

    return doc

def converti(docx_path, output_dir, soffice, fmt):
    filtro = "pdf:writer_pdf_Export:EmbedStandardFonts=true,SelectPdfVersion=1" if fmt == "pdfa" else "pdf"
    subprocess.run([soffice, "--headless", "--convert-to", filtro, "--outdir", output_dir, docx_path],
                   check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

# ── Finestra selezione record ─────────────────────────────────────────────────
class DialogSelezioneRecord(tk.Toplevel):
    def __init__(self, parent, intestazioni, righe, selezionati):
        super().__init__(parent)
        self.title("Selezione record")
        self.geometry("860x540")
        self.resizable(True, True)
        self.configure(bg=BG)
        self.grab_set()
        self.intestazioni = intestazioni
        self.righe = righe
        self.risultato = None
        self._checked = {i: (i in selezionati) for i in range(len(righe))}
        self._item_to_idx = {}
        self._build()
        self.protocol("WM_DELETE_WINDOW", self._annulla)

    def _build(self):
        toolbar = tk.Frame(self, bg=ACCENT, height=48)
        toolbar.pack(fill="x")
        toolbar.pack_propagate(False)
        tk.Label(toolbar, text="Seleziona i record da elaborare",
                 font=(FONT, 11, "bold"), bg=ACCENT, fg="white").pack(side="left", padx=16)
        self._lbl_count = tk.StringVar()
        tk.Label(toolbar, textvariable=self._lbl_count, font=(FONT, 9),
                 bg=ACCENT, fg="#C7D2FE").pack(side="right", padx=16)

        ctrl = tk.Frame(self, bg=BG, pady=8)
        ctrl.pack(fill="x", padx=16)
        tk.Label(ctrl, text="🔍 Cerca:", font=(FONT, 9), bg=BG, fg=TEXT).pack(side="left")
        self._search_var = tk.StringVar()
        self._search_var.trace_add("write", lambda *_: self._filtra())
        tk.Entry(ctrl, textvariable=self._search_var, font=(FONT, 9),
                 bg=CARD, relief="flat", highlightbackground=BORDER,
                 highlightthickness=1, width=28).pack(side="left", padx=(4, 16), ipady=4)
        for txt, cmd in [("Seleziona tutti", self._sel_tutti),
                          ("Deseleziona tutti", self._desel_tutti),
                          ("Inverti", self._inverti)]:
            tk.Button(ctrl, text=txt, font=(FONT, 9), bg=CARD, fg=ACCENT,
                      relief="flat", cursor="hand2", padx=8,
                      highlightbackground=BORDER, highlightthickness=1,
                      command=cmd).pack(side="left", padx=4, ipady=4)

        table_frame = tk.Frame(self, bg=CARD, highlightbackground=BORDER, highlightthickness=1)
        table_frame.pack(fill="both", expand=True, padx=16, pady=(0, 8))
        cols = ("sel",) + tuple(self.intestazioni)
        self.tree = ttk.Treeview(table_frame, columns=cols, show="headings", selectmode="none")
        style = ttk.Style()
        style.configure("Treeview", font=(FONT, 9), rowheight=26, background=CARD, fieldbackground=CARD)
        style.configure("Treeview.Heading", font=(FONT, 9, "bold"))
        self.tree.heading("sel", text="✔")
        self.tree.column("sel", width=36, anchor="center", stretch=False)
        for col in self.intestazioni:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=max(80, min(180, 700 // max(len(self.intestazioni), 1))), anchor="w")
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")
        self.tree.pack(fill="both", expand=True)
        self.tree.bind("<Button-1>", self._toggle_click)
        self._popola()

        footer = tk.Frame(self, bg=BG)
        footer.pack(fill="x", padx=16, pady=(0, 12))
        tk.Button(footer, text="Annulla", font=(FONT, 10), bg=CARD, fg=TEXT,
                  relief="flat", cursor="hand2", padx=16,
                  highlightbackground=BORDER, highlightthickness=1,
                  command=self._annulla).pack(side="right", padx=(8, 0), ipady=6)
        tk.Button(footer, text="✔  Conferma", font=(FONT, 10, "bold"),
                  bg=ACCENT, fg="white", relief="flat", cursor="hand2", padx=16,
                  command=self._conferma).pack(side="right", ipady=6)
        self._aggiorna_contatore()

    def _popola(self):
        self.tree.delete(*self.tree.get_children())
        self._item_to_idx.clear()
        filtro = self._search_var.get().lower() if hasattr(self, '_search_var') else ""
        for i, riga in enumerate(self.righe):
            valori = [str(riga.get(col, "")) for col in self.intestazioni]
            if filtro and not any(filtro in v.lower() for v in valori): continue
            chk = "☑" if self._checked.get(i, True) else "☐"
            tag = "checked" if self._checked.get(i, True) else "unchecked"
            iid = self.tree.insert("", "end", values=(chk,) + tuple(valori), tags=(tag,))
            self._item_to_idx[iid] = i
        self.tree.tag_configure("checked", background=CARD)
        self.tree.tag_configure("unchecked", background="#F9FAFB", foreground=MUTED)
        self._aggiorna_contatore()

    def _filtra(self): self._popola()

    def _toggle_click(self, e):
        iid = self.tree.identify_row(e.y)
        if not iid: return
        idx = self._item_to_idx.get(iid)
        if idx is None: return
        self._checked[idx] = not self._checked[idx]
        chk = "☑" if self._checked[idx] else "☐"
        tag = "checked" if self._checked[idx] else "unchecked"
        vals = self.tree.item(iid, "values")
        self.tree.item(iid, values=(chk,) + tuple(vals[1:]), tags=(tag,))
        self.tree.tag_configure("checked", background=CARD)
        self.tree.tag_configure("unchecked", background="#F9FAFB", foreground=MUTED)
        self._aggiorna_contatore()

    def _sel_tutti(self):
        for iid, idx in self._item_to_idx.items():
            self._checked[idx] = True
            vals = self.tree.item(iid, "values")
            self.tree.item(iid, values=("☑",) + tuple(vals[1:]), tags=("checked",))
        self.tree.tag_configure("checked", background=CARD)
        self._aggiorna_contatore()

    def _desel_tutti(self):
        for iid, idx in self._item_to_idx.items():
            self._checked[idx] = False
            vals = self.tree.item(iid, "values")
            self.tree.item(iid, values=("☐",) + tuple(vals[1:]), tags=("unchecked",))
        self.tree.tag_configure("unchecked", background="#F9FAFB", foreground=MUTED)
        self._aggiorna_contatore()

    def _inverti(self):
        for iid, idx in self._item_to_idx.items():
            self._checked[idx] = not self._checked[idx]
            chk = "☑" if self._checked[idx] else "☐"
            tag = "checked" if self._checked[idx] else "unchecked"
            vals = self.tree.item(iid, "values")
            self.tree.item(iid, values=(chk,) + tuple(vals[1:]), tags=(tag,))
        self.tree.tag_configure("checked", background=CARD)
        self.tree.tag_configure("unchecked", background="#F9FAFB", foreground=MUTED)
        self._aggiorna_contatore()

    def _aggiorna_contatore(self):
        sel = sum(1 for v in self._checked.values() if v)
        self._lbl_count.set(f"{sel} di {len(self.righe)} selezionati")

    def _conferma(self):
        self.risultato = [i for i, v in self._checked.items() if v]
        self.destroy()

    def _annulla(self):
        self.risultato = None
        self.destroy()

# ── App principale ────────────────────────────────────────────────────────────
class MailMergeApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Mail Merge")
        self.geometry("720x780")
        self.minsize(640, 600)
        self.resizable(True, True)
        self.configure(bg=BG)

        self.xlsx_path = tk.StringVar()
        self.docx_path = tk.StringVar()
        self.out_dir   = tk.StringVar()
        self.fmt       = tk.StringVar(value="pdf")
        self.running   = False
        self._stop_flag = False
        self._intestazioni = []
        self._righe = []
        self._righe_selezionate = []

        self.nome_fisso = tk.StringVar(value="documento")
        self.sep1       = tk.StringVar(value="_")
        self.sep2       = tk.StringVar(value="_")
        self.campo1     = tk.StringVar(value="(nessuno)")
        self.campo2     = tk.StringVar(value="(nessuno)")
        # Ordine blocchi: "fisso", "campo1", "campo2"
        self._blocchi_ordine = ["fisso", "campo1", "campo2"]
        # Colonne con formato valuta
        self._colonne_valuta = set()

        self._build_ui()

        for var in (self.nome_fisso, self.sep1, self.campo1, self.sep2, self.campo2, self.fmt):
            var.trace_add("write", lambda *_: self._aggiorna_anteprima_nome())

    def _build_ui(self):
        # ── Header fisso ──
        hdr = tk.Frame(self, bg=ACCENT, height=56)
        hdr.pack(fill="x", side="top")
        hdr.pack_propagate(False)
        tk.Label(hdr, text="📄  Mail Merge", font=(FONT, 15, "bold"),
                 bg=ACCENT, fg="white").pack(side="left", padx=20)
        tk.Label(hdr, text="Genera documenti personalizzati da template Word",
                 font=(FONT, 9), bg=ACCENT, fg="#C7D2FE").pack(side="left")

        # ── Bottone fisso in fondo ──
        btn_frame = tk.Frame(self, bg=BG)
        btn_frame.pack(fill="x", side="bottom", padx=20, pady=(0, 16))
        self.btn = tk.Button(btn_frame, text="▶   Avvia generazione",
                             font=(FONT, 11, "bold"), bg=ACCENT, fg="white",
                             activebackground=ACCENT2, activeforeground="white",
                             relief="flat", cursor="hand2", height=2,
                             command=self._avvia)
        self.btn.pack(fill="x")

        # ── Area scrollabile ──
        scroll_container = tk.Frame(self, bg=BG)
        scroll_container.pack(fill="both", expand=True, side="top")

        vscroll = ttk.Scrollbar(scroll_container, orient="vertical")
        vscroll.pack(side="right", fill="y")
        hscroll = ttk.Scrollbar(scroll_container, orient="horizontal")
        hscroll.pack(side="bottom", fill="x")

        canvas = tk.Canvas(scroll_container, bg=BG, highlightthickness=0,
                           yscrollcommand=vscroll.set, xscrollcommand=hscroll.set)
        canvas.pack(side="left", fill="both", expand=True)
        vscroll.config(command=canvas.yview)
        hscroll.config(command=canvas.xview)

        self._body = tk.Frame(canvas, bg=BG)
        body_id = canvas.create_window((0, 0), window=self._body, anchor="nw")

        def on_configure(e):
            canvas.configure(scrollregion=canvas.bbox("all"))
        def on_canvas_resize(e):
            canvas.itemconfig(body_id, width=e.width)
        self._body.bind("<Configure>", on_configure)
        canvas.bind("<Configure>", on_canvas_resize)

        # Scroll con rotella del mouse
        def on_mousewheel(e):
            canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
        canvas.bind_all("<MouseWheel>", on_mousewheel)

        self._build_body(self._body)

    def _build_body(self, body):
        pad = {"padx": 20, "pady": (0, 12)}

        # ── Origine dati ──
        xlsx_sec = self._section(body, "📊  Origine dati (Excel)")
        xlsx_inner = tk.Frame(xlsx_sec, bg=CARD)
        xlsx_inner.pack(fill="x", padx=16, pady=10)
        self._xlsx_entry = tk.Entry(xlsx_inner, textvariable=self.xlsx_path,
                                    font=(FONT, 9), fg=MUTED, bg="#F9FAFB",
                                    relief="flat", highlightbackground=BORDER, highlightthickness=1)
        self._xlsx_entry.insert(0, "Seleziona file .xlsx")
        self._xlsx_entry.pack(side="left", fill="x", expand=True, ipady=6, padx=(0, 8))
        tk.Button(xlsx_inner, text="Sfoglia…", font=(FONT, 9), bg=ACCENT, fg="white",
                  activebackground=ACCENT2, activeforeground="white",
                  relief="flat", cursor="hand2", padx=12,
                  command=self._carica_xlsx).pack(side="left", ipady=5, padx=(0, 8))
        self.btn_record = tk.Button(xlsx_inner, text="👁  Vedi record",
                                    font=(FONT, 9), bg=CARD, fg=ACCENT,
                                    relief="flat", cursor="hand2", padx=8,
                                    highlightbackground=BORDER, highlightthickness=1,
                                    state="disabled", command=self._apri_selezione)
        self.btn_record.pack(side="left", ipady=5)
        self._lbl_record = tk.Label(xlsx_sec, text="", font=(FONT, 9),
                                    bg=CARD, fg=MUTED, padx=16)
        self._lbl_record.pack(anchor="w", pady=(0, 8))

        # ── Template Word ──
        self._card(body, "📝  Template documento (Word)", self.docx_path,
                   "Seleziona file .docx",
                   lambda: self._browse_file(self.docx_path, [("Word", "*.docx")]))

        # ── Output ──
        self._card(body, "📁  Cartella di output", self.out_dir,
                   "Seleziona cartella di destinazione",
                   lambda: self._browse_dir(self.out_dir))

        # ── Nome file ──
        nome_sec = self._section(body, "🏷️  Nome file esportato")
        self._nome_container = tk.Frame(nome_sec, bg=CARD)
        self._nome_container.pack(fill="x", padx=16, pady=(8, 4))

        sep_vals = ["_", "-", " ", ".", ""]

        # Blocco parte fissa
        self._frame_fisso = tk.Frame(self._nome_container, bg=CARD,
                                      highlightbackground=BORDER, highlightthickness=1)
        self._frame_fisso.pack(fill="x", pady=3)
        self._build_blocco_fisso(self._frame_fisso)

        # Blocco campo1
        self._frame_campo1 = tk.Frame(self._nome_container, bg=CARD,
                                       highlightbackground=BORDER, highlightthickness=1)
        self._frame_campo1.pack(fill="x", pady=3)
        self._build_blocco_campo(self._frame_campo1, "campo1", self.sep1, self.campo1)

        # Blocco campo2
        self._frame_campo2 = tk.Frame(self._nome_container, bg=CARD,
                                       highlightbackground=BORDER, highlightthickness=1)
        self._frame_campo2.pack(fill="x", pady=3)
        self._build_blocco_campo(self._frame_campo2, "campo2", self.sep2, self.campo2)

        # Anteprima
        prev_row = tk.Frame(nome_sec, bg=CARD)
        prev_row.pack(fill="x", padx=16, pady=(4, 10))
        tk.Label(prev_row, text="Anteprima:", font=(FONT, 8), bg=CARD, fg=MUTED).pack(side="left")
        self._lbl_anteprima = tk.Label(prev_row, text="documento.pdf",
                                        font=("Consolas", 9), bg="#EEF2FF",
                                        fg=ACCENT, padx=8, pady=3)
        self._lbl_anteprima.pack(side="left", padx=(6, 0))

        # ── Formato valuta ──
        valuta_sec = self._section(body, "💶  Formato valuta")
        valuta_inner = tk.Frame(valuta_sec, bg=CARD)
        valuta_inner.pack(fill="x", padx=16, pady=(6, 12))
        self._lbl_valuta_info = tk.Label(valuta_inner,
                                          text="Nessuna colonna configurata come valuta.",
                                          font=(FONT, 9), bg=CARD, fg=MUTED)
        self._lbl_valuta_info.pack(side="left", expand=True, anchor="w")
        self._btn_valuta = tk.Button(valuta_inner, text="⚙️  Configura colonne",
                                      font=(FONT, 9), bg=ACCENT, fg="white",
                                      activebackground=ACCENT2, activeforeground="white",
                                      relief="flat", cursor="hand2", padx=12,
                                      state="disabled",
                                      command=self._apri_dialog_valuta)
        self._btn_valuta.pack(side="right", ipady=5)

        # ── Formato ──
        fmt_sec = self._section(body, "📤  Formato di esportazione")
        fmt_inner = tk.Frame(fmt_sec, bg=CARD)
        fmt_inner.pack(fill="x", padx=16, pady=(4, 14))
        for val, lbl, desc in [("pdf",  "PDF",   "Documento PDF standard"),
                                ("pdfa", "PDF/A", "PDF archiviabile a lungo termine"),
                                ("docx", "Word",  "Documento Word modificabile")]:
            row = tk.Frame(fmt_inner, bg=CARD)
            row.pack(fill="x", pady=4)
            tk.Radiobutton(row, text=lbl, variable=self.fmt, value=val,
                           font=(FONT, 10, "bold"), bg=CARD, fg=TEXT,
                           activebackground=CARD, selectcolor=CARD).pack(side="left")
            tk.Label(row, text=desc, font=(FONT, 9), bg=CARD,
                     fg=MUTED).pack(side="left", padx=8)

        # ── Progress + Log ──
        # ── Progress + Log ──
        prog_frame = tk.Frame(body, bg=BG)
        prog_frame.pack(fill="x", padx=20, pady=(4, 0))

        # Riga contatore e pulsante stop
        top_row = tk.Frame(prog_frame, bg=BG)
        top_row.pack(fill="x", pady=(0, 4))
        self.prog_label = tk.Label(top_row, text="", font=(FONT, 9), bg=BG, fg=MUTED)
        self.prog_label.pack(side="left")
        self.btn_stop = tk.Button(top_row, text="⏹  Interrompi",
                                   font=(FONT, 9, "bold"), bg="#DC2626", fg="white",
                                   activebackground="#B91C1C", activeforeground="white",
                                   relief="flat", cursor="hand2", padx=10,
                                   state="disabled", command=self._stop)
        self.btn_stop.pack(side="right", ipady=3)

        # Barra di avanzamento
        self.prog_bar = ttk.Progressbar(prog_frame, mode="determinate")
        self.prog_bar.pack(fill="x", pady=(0, 4))

        # Etichetta dettaglio sotto la barra
        self.prog_detail = tk.Label(prog_frame, text="", font=("Consolas", 8),
                                     bg=BG, fg=MUTED)
        self.prog_detail.pack(anchor="w")

        log_frame = tk.Frame(body, bg=CARD, highlightbackground=BORDER, highlightthickness=1)
        log_frame.pack(fill="x", padx=20, pady=(12, 20))
        self.log = tk.Text(log_frame, height=7, font=("Consolas", 9),
                           bg=CARD, fg=TEXT, relief="flat", state="disabled",
                           wrap="word", padx=10, pady=8)
        self.log.pack(fill="both")
        self.log.tag_config("ok",   foreground=SUCCESS)
        self.log.tag_config("err",  foreground=ERROR)
        self.log.tag_config("info", foreground=ACCENT)

    def _build_blocco_fisso(self, frame):
        inner = tk.Frame(frame, bg="#F0F0FF", padx=8, pady=6)
        inner.pack(fill="x")
        tk.Label(inner, text="📌 Parte fissa", font=(FONT, 8, "bold"),
                 bg="#F0F0FF", fg=ACCENT, width=14, anchor="w").pack(side="left")
        tk.Entry(inner, textvariable=self.nome_fisso, font=(FONT, 9),
                 bg=CARD, relief="flat", highlightbackground=BORDER,
                 highlightthickness=1, width=22).pack(side="left", ipady=4, padx=(4, 12))
        tk.Button(inner, text="▲", font=(FONT, 8), bg=CARD, fg=TEXT,
                  relief="flat", cursor="hand2", padx=4,
                  command=lambda: self._sposta("fisso", -1)).pack(side="right", padx=2)
        tk.Button(inner, text="▼", font=(FONT, 8), bg=CARD, fg=TEXT,
                  relief="flat", cursor="hand2", padx=4,
                  command=lambda: self._sposta("fisso", 1)).pack(side="right", padx=2)
        tk.Label(inner, text="Ordine:", font=(FONT, 8), bg="#F0F0FF", fg=MUTED).pack(side="right", padx=4)

    def _build_blocco_campo(self, frame, chiave, sep_var, campo_var):
        etichetta = "🔵 Campo dinamico 1" if chiave == "campo1" else "🟢 Campo dinamico 2"
        colore_bg = "#F0FFF4" if chiave == "campo2" else "#EFF6FF"
        inner = tk.Frame(frame, bg=colore_bg, padx=8, pady=6)
        inner.pack(fill="x")
        tk.Label(inner, text=etichetta, font=(FONT, 8, "bold"),
                 bg=colore_bg, fg=TEXT, width=18, anchor="w").pack(side="left")
        tk.Label(inner, text="Sep.", font=(FONT, 8), bg=colore_bg, fg=MUTED).pack(side="left", padx=(0, 2))
        ttk.Combobox(inner, textvariable=sep_var,
                     values=["_", "-", " ", ".", ""],
                     width=4, font=(FONT, 9), state="readonly").pack(side="left", padx=(0, 8))
        tk.Label(inner, text="Colonna:", font=(FONT, 8), bg=colore_bg, fg=MUTED).pack(side="left", padx=(0, 2))
        cb = ttk.Combobox(inner, textvariable=campo_var,
                          values=["(nessuno)"] + self._intestazioni,
                          width=20, font=(FONT, 9), state="readonly")
        cb.pack(side="left", padx=(0, 12))
        if chiave == "campo1":
            self._cb_campo1 = cb
        else:
            self._cb_campo2 = cb
        tk.Button(inner, text="▲", font=(FONT, 8), bg=CARD, fg=TEXT,
                  relief="flat", cursor="hand2", padx=4,
                  command=lambda k=chiave: self._sposta(k, -1)).pack(side="right", padx=2)
        tk.Button(inner, text="▼", font=(FONT, 8), bg=CARD, fg=TEXT,
                  relief="flat", cursor="hand2", padx=4,
                  command=lambda k=chiave: self._sposta(k, 1)).pack(side="right", padx=2)
        tk.Label(inner, text="Ordine:", font=(FONT, 8), bg=colore_bg, fg=MUTED).pack(side="right", padx=4)

    def _sposta(self, chiave, direzione):
        idx = self._blocchi_ordine.index(chiave)
        nuovo = idx + direzione
        if nuovo < 0 or nuovo >= len(self._blocchi_ordine): return
        self._blocchi_ordine[idx], self._blocchi_ordine[nuovo] = \
            self._blocchi_ordine[nuovo], self._blocchi_ordine[idx]
        self._ridisegna_blocchi()
        self._aggiorna_anteprima_nome()

    def _ridisegna_blocchi(self):
        mappa = {
            "fisso":  self._frame_fisso,
            "campo1": self._frame_campo1,
            "campo2": self._frame_campo2,
        }
        for w in self._nome_container.winfo_children():
            w.pack_forget()
        for chiave in self._blocchi_ordine:
            mappa[chiave].pack(fill="x", pady=3)

    def _apri_dialog_valuta(self):
        dlg = tk.Toplevel(self)
        dlg.title("Configura formato valuta")
        dlg.geometry("380x460")
        dlg.resizable(False, True)
        dlg.configure(bg=BG)
        dlg.grab_set()

        # Header
        hdr = tk.Frame(dlg, bg=ACCENT, height=44)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        tk.Label(hdr, text="💶  Seleziona colonne da formattare come valuta",
                 font=(FONT, 9, "bold"), bg=ACCENT, fg="white").pack(side="left", padx=12)

        tk.Label(dlg, text="Le colonne selezionate verranno esportate nel formato  € 1.234,56",
                 font=(FONT, 8), bg=BG, fg=MUTED, wraplength=340).pack(padx=16, pady=(10, 6), anchor="w")

        # Lista scrollabile
        list_frame = tk.Frame(dlg, bg=CARD, highlightbackground=BORDER, highlightthickness=1)
        list_frame.pack(fill="both", expand=True, padx=16, pady=(0, 8))

        canvas = tk.Canvas(list_frame, bg=CARD, highlightthickness=0)
        vsb = ttk.Scrollbar(list_frame, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        inner = tk.Frame(canvas, bg=CARD)
        win_id = canvas.create_window((0, 0), window=inner, anchor="nw")

        def on_cfg(e): canvas.configure(scrollregion=canvas.bbox("all"))
        def on_resize(e): canvas.itemconfig(win_id, width=e.width)
        inner.bind("<Configure>", on_cfg)
        canvas.bind("<Configure>", on_resize)
        canvas.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        # Checkbox per ogni colonna
        checks = {}
        for col in self._intestazioni:
            var = tk.BooleanVar(value=(col in self._colonne_valuta))
            checks[col] = var
            row = tk.Frame(inner, bg=CARD)
            row.pack(fill="x", padx=8, pady=2)
            tk.Checkbutton(row, text=col, variable=var,
                           font=(FONT, 9), bg=CARD, fg=TEXT,
                           activebackground=CARD, selectcolor=CARD).pack(side="left")
            # Mostra esempio valore dal primo record
            if self._righe:
                val = self._righe[0].get(col, "")
                try:
                    n = float(str(val).replace(",", "."))
                    esempio = f"→  € {n:,.2f}".replace(",","X").replace(".",",").replace("X",".")
                    tk.Label(row, text=esempio, font=("Consolas", 8),
                             bg=CARD, fg=SUCCESS).pack(side="right", padx=8)
                except (ValueError, TypeError):
                    pass

        # Footer
        footer = tk.Frame(dlg, bg=BG)
        footer.pack(fill="x", padx=16, pady=(0, 12))

        def seleziona_tutti():
            for v in checks.values(): v.set(True)
        def deseleziona_tutti():
            for v in checks.values(): v.set(False)

        tk.Button(footer, text="Seleziona tutti", font=(FONT, 9), bg=CARD, fg=ACCENT,
                  relief="flat", cursor="hand2", padx=8,
                  highlightbackground=BORDER, highlightthickness=1,
                  command=seleziona_tutti).pack(side="left", ipady=4, padx=(0, 6))
        tk.Button(footer, text="Deseleziona tutti", font=(FONT, 9), bg=CARD, fg=ACCENT,
                  relief="flat", cursor="hand2", padx=8,
                  highlightbackground=BORDER, highlightthickness=1,
                  command=deseleziona_tutti).pack(side="left", ipady=4)

        def conferma():
            self._colonne_valuta = {col for col, v in checks.items() if v.get()}
            self._aggiorna_label_valuta()
            dlg.destroy()

        tk.Button(footer, text="✔  Conferma", font=(FONT, 10, "bold"),
                  bg=ACCENT, fg="white", activebackground=ACCENT2,
                  relief="flat", cursor="hand2", padx=16,
                  command=conferma).pack(side="right", ipady=6)
        tk.Button(footer, text="Annulla", font=(FONT, 9), bg=CARD, fg=TEXT,
                  relief="flat", cursor="hand2", padx=12,
                  highlightbackground=BORDER, highlightthickness=1,
                  command=dlg.destroy).pack(side="right", ipady=4, padx=(0, 8))

    def _aggiorna_label_valuta(self):
        if not self._colonne_valuta:
            self._lbl_valuta_info.config(text="Nessuna colonna configurata come valuta.", fg=MUTED)
        else:
            cols = ", ".join(sorted(self._colonne_valuta))
            self._lbl_valuta_info.config(text=f"Valuta: {cols}", fg=SUCCESS)

    def _aggiorna_sezione_valuta(self):
        pass  # mantenuto per compatibilità

    def _aggiorna_colonne_valuta(self):
        pass  # mantenuto per compatibilità

    def _section(self, parent, title):
        frame = tk.Frame(parent, bg=CARD, highlightbackground=BORDER, highlightthickness=1)
        frame.pack(fill="x", padx=20, pady=(0, 12))
        tk.Label(frame, text=title, font=(FONT, 10, "bold"),
                 bg=CARD, fg=TEXT, padx=16, pady=10).pack(anchor="w")
        tk.Frame(frame, bg=BORDER, height=1).pack(fill="x")
        return frame

    def _card(self, parent, title, var, placeholder, cmd):
        frame = self._section(parent, title)
        inner = tk.Frame(frame, bg=CARD)
        inner.pack(fill="x", padx=16, pady=10)
        e = tk.Entry(inner, textvariable=var, font=(FONT, 9), fg=MUTED,
                     bg="#F9FAFB", relief="flat",
                     highlightbackground=BORDER, highlightthickness=1)
        e.insert(0, placeholder)
        e.pack(side="left", fill="x", expand=True, ipady=6, padx=(0, 10))
        def fi(ev, en=e, ph=placeholder):
            if en.get() == ph: en.delete(0, "end"); en.config(fg=TEXT)
        def fo(ev, en=e, ph=placeholder):
            if not en.get(): en.insert(0, ph); en.config(fg=MUTED)
        e.bind("<FocusIn>", fi)
        e.bind("<FocusOut>", fo)
        tk.Button(inner, text="Sfoglia…", font=(FONT, 9), bg=ACCENT, fg="white",
                  activebackground=ACCENT2, activeforeground="white",
                  relief="flat", cursor="hand2", padx=12,
                  command=cmd).pack(side="right", ipady=5)

    def _browse_file(self, var, types):
        p = filedialog.askopenfilename(filetypes=types)
        if p: var.set(p)

    def _browse_dir(self, var):
        p = filedialog.askdirectory()
        if p: var.set(p)

    def _carica_xlsx(self):
        path = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx")])
        if not path: return
        try:
            wb = openpyxl.load_workbook(path)
            ws = wb.active
            self._intestazioni = [str(c.value).strip() if c.value else ""
                                   for c in next(ws.iter_rows(max_row=1))]
            self._righe = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                if any(v is not None for v in row):
                    self._righe.append({self._intestazioni[i]: (str(v) if v is not None else "")
                                        for i, v in enumerate(row)})
            self._righe_selezionate = list(range(len(self._righe)))
            self.xlsx_path.set(path)
            self._xlsx_entry.config(fg=TEXT)
            self.btn_record.config(state="normal")
            self._aggiorna_label_record()
            # Aggiorna combobox campi dinamici
            cols = ["(nessuno)"] + self._intestazioni
            self._cb_campo1.config(values=cols)
            self._cb_campo2.config(values=cols)
            if self.campo1.get() not in cols: self.campo1.set("(nessuno)")
            if self.campo2.get() not in cols: self.campo2.set("(nessuno)")
            self._btn_valuta.config(state="normal")
            self._colonne_valuta = set()
            self._valuta_checks = {}
            self._aggiorna_label_valuta()
            self._aggiorna_anteprima_nome()
            self._log(f"Excel caricato: {len(self._righe)} record, colonne: {', '.join(self._intestazioni)}", "info")
        except Exception as e:
            messagebox.showerror("Errore", f"Impossibile leggere il file Excel:\n{e}")

    def _apri_selezione(self):
        dlg = DialogSelezioneRecord(self, self._intestazioni, self._righe, self._righe_selezionate)
        self.wait_window(dlg)
        if dlg.risultato is not None:
            self._righe_selezionate = dlg.risultato
            self._aggiorna_label_record()
            self._aggiorna_anteprima_nome()

    def _aggiorna_label_record(self):
        tot = len(self._righe)
        sel = len(self._righe_selezionate)
        self._lbl_record.config(
            text=f"✔  {sel} di {tot} record selezionati",
            fg=SUCCESS if sel > 0 else ERROR)

    def _formatta_valore(self, colonna, valore):
        """Formatta il valore come valuta se la colonna è marcata."""
        if colonna in self._colonne_valuta:
            try:
                n = float(str(valore).replace(",", "."))
                return f"€ {n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            except (ValueError, TypeError):
                return valore
        return valore

    def _get_parti_nome(self, riga):
        """Restituisce le parti del nome file nell'ordine corrente."""
        esempio = riga or {}
        parti = []
        mappa_val = {
            "fisso":  (self.nome_fisso.get().strip(), None, None),
            "campo1": (None, self.sep1, self.campo1),
            "campo2": (None, self.sep2, self.campo2),
        }
        for chiave in self._blocchi_ordine:
            fisso_val, sep_var, campo_var = mappa_val[chiave]
            if chiave == "fisso":
                if fisso_val: parti.append(fisso_val)
            else:
                c = campo_var.get()
                if c and c != "(nessuno)":
                    val = str(esempio.get(c, c))
                    if parti: parti.append(sep_var.get())
                    parti.append(val)
        return parti

    def _aggiorna_anteprima_nome(self):
        esempio = self._righe[self._righe_selezionate[0]] if self._righe and self._righe_selezionate else {}
        parti = self._get_parti_nome(esempio)
        nome = re.sub(r'[^a-zA-Z0-9_\-. ]', '_', "".join(parti) or "documento")
        ext = "pdf" if self.fmt.get() in ("pdf", "pdfa") else "docx"
        self._lbl_anteprima.config(text=f"{nome}.{ext}")

    def _build_nome_file(self, riga):
        parti = self._get_parti_nome(riga)
        return re.sub(r'[^a-zA-Z0-9_\-.]', '_', "".join(parti) or "documento")

    def _log(self, msg, tag=""):
        self.log.config(state="normal")
        self.log.insert("end", msg + "\n", tag)
        self.log.see("end")
        self.log.config(state="disabled")

    def _avvia(self):
        if self.running: return
        xlsx = self.xlsx_path.get()
        docx = self.docx_path.get()
        out  = self.out_dir.get()
        fmt  = self.fmt.get()
        if not os.path.isfile(xlsx) or not xlsx.endswith(".xlsx"):
            messagebox.showerror("Errore", "Seleziona un file Excel (.xlsx) valido."); return
        if not os.path.isfile(docx) or not docx.endswith(".docx"):
            messagebox.showerror("Errore", "Seleziona un file Word (.docx) valido."); return
        if not os.path.isdir(out):
            messagebox.showerror("Errore", "Seleziona una cartella di output valida."); return
        if not self._righe_selezionate:
            messagebox.showwarning("Nessun record", "Nessun record selezionato."); return
        if fmt in ("pdf", "pdfa"):
            soffice = trova_libreoffice()
            if not soffice:
                if messagebox.askyesno("LibreOffice non trovato",
                    "LibreOffice è necessario per i PDF.\nVuoi installarlo adesso?"):
                    threading.Thread(target=self._installa_libreoffice, daemon=True).start()
                return
        else:
            soffice = None
        self._stop_flag = False
        self.running = True
        self.btn.config(state="disabled", text="⏳  Generazione in corso...")
        self.btn_stop.config(state="normal")
        self.log.config(state="normal"); self.log.delete("1.0", "end"); self.log.config(state="disabled")
        self.prog_bar["value"] = 0
        self.prog_label.config(text="")
        self.prog_detail.config(text="")
        righe = [self._righe[i] for i in self._righe_selezionate]
        threading.Thread(target=self._run, args=(docx, out, fmt, soffice, righe), daemon=True).start()

    def _stop(self):
        self._stop_flag = True
        self.btn_stop.config(state="disabled", text="⏹  Interruzione...")
        self._log("⚠️  Interruzione richiesta — attendi il completamento del file corrente...", "err")

    def _run(self, docx, out, fmt, soffice, righe):
        try:
            totale = len(righe)
            self._log(f"Avvio elaborazione di {totale} record...", "info")
            tmp = os.path.join(out, "_tmp_merge")
            os.makedirs(tmp, exist_ok=True)
            ok = err = 0
            for idx, riga in enumerate(righe):
                if self._stop_flag:
                    self._log(f"\n⏹  Generazione interrotta dopo {idx} di {totale} record.", "err")
                    break
                fname = self._build_nome_file(riga)
                docx_tmp = os.path.join(tmp, f"{fname}.docx")
                try:
                    doc = processa_documento(docx, riga, self._colonne_valuta)
                    doc.save(docx_tmp)
                    if fmt in ("pdf", "pdfa"):
                        converti(docx_tmp, out, soffice, fmt)
                    else:
                        shutil.copy(docx_tmp, os.path.join(out, f"{fname}.docx"))
                    ext = "pdf" if fmt in ("pdf", "pdfa") else "docx"
                    self._log(f"  [{idx+1}/{totale}]  {fname}.{ext}", "ok")
                    ok += 1
                except Exception as e:
                    self._log(f"  [{idx+1}/{totale}]  ERRORE su {fname}: {e}", "err")
                    err += 1

                pct = int(((idx+1) / totale) * 100)
                self.prog_bar["value"] = pct
                self.prog_label.config(
                    text=f"Elaborati {idx+1} di {totale}  ({pct}%)")
                self.prog_detail.config(
                    text=f"✔ {ok} completati   ✖ {err} errori   "
                         f"{'⏹ interrotto' if self._stop_flag else f'▶ rimanenti: {totale - idx - 1}'}")
                self.update_idletasks()

            shutil.rmtree(tmp, ignore_errors=True)
            if self._stop_flag:
                self.prog_label.config(text=f"Interrotto: {ok} generati su {totale}, {err} errori.")
            else:
                self.prog_label.config(text=f"Completato: {ok} generati, {err} errori.")
                self.prog_detail.config(text=f"✔ {ok} completati   ✖ {err} errori   ▶ rimanenti: 0")
            self._log(f"\n{'⏹' if self._stop_flag else '✅'}  {'Interrotto' if self._stop_flag else 'Completato'}: {ok} file in '{out}'", "info")
            if err: self._log(f"⚠️  {err} errori riscontrati.", "err")
            if ok and not self._stop_flag:
                messagebox.showinfo("Completato", f"{ok} file generati con successo in:\n{out}")
            elif ok and self._stop_flag:
                messagebox.showwarning("Interrotto", f"Generazione interrotta.\n{ok} file salvati in:\n{out}")
        except Exception as e:
            self._log(f"Errore generale: {e}", "err")
            messagebox.showerror("Errore", str(e))
        self.running = False
        self._stop_flag = False
        self.btn.config(state="normal", text="▶   Avvia generazione")
        self.btn_stop.config(state="disabled", text="⏹  Interrompi")

    def _installa_libreoffice(self):
        self.btn.config(state="disabled", text="⏳  Installazione LibreOffice...")
        self.prog_label.config(text="Installazione in corso...")
        self.prog_bar.config(mode="indeterminate")
        self.prog_bar.start(10)
        self._log("Installazione LibreOffice tramite winget...", "info")
        try:
            r = subprocess.run([
                "winget", "install", "--id", "TheDocumentFoundation.LibreOffice",
                "--silent", "--accept-package-agreements", "--accept-source-agreements"
            ], capture_output=True, text=True)
            self.prog_bar.stop()
            self.prog_bar.config(mode="determinate")
            if r.returncode == 0 and trova_libreoffice():
                self._log("LibreOffice installato!", "ok")
                self.prog_label.config(text="LibreOffice installato. Pronto.")
                messagebox.showinfo("Installazione completata", "LibreOffice installato!\nOra puoi generare i PDF.")
            else:
                messagebox.showinfo("Riavvio necessario",
                    "LibreOffice installato.\nChiudi e riavvia l'applicazione.")
                self.destroy()
        except FileNotFoundError:
            self.prog_bar.stop(); self.prog_bar.config(mode="determinate")
            self._log("winget non disponibile.", "err")
            messagebox.showerror("Errore", "Scarica LibreOffice da:\nhttps://www.libreoffice.org/download/")
        except Exception as e:
            self.prog_bar.stop(); self.prog_bar.config(mode="determinate")
            self._log(f"Errore: {e}", "err")
        finally:
            self.btn.config(state="normal", text="▶   Avvia generazione")
            self.prog_label.config(text="")

if __name__ == "__main__":
    app = MailMergeApp()
    app.mainloop()